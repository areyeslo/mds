import json, os, docx, io, base64

from flask import current_app
from sqlalchemy.schema import FetchedValue
from sqlalchemy.dialects.postgresql import UUID
from sqlalchemy.ext.hybrid import hybrid_property
from docx.shared import Inches
from datetime import datetime

from app.extensions import db
from app.api.utils.models_mixins import AuditMixin, Base
from bs4 import BeautifulSoup


def get_model_by_model_name(model_name):
    for c in Base._decl_class_registry.values():
        if hasattr(c, '__name__') and c.__name__ == model_name:
            return c


class DocumentTemplate(Base, AuditMixin):
    __tablename__ = 'document_template'

    document_template_code = db.Column(db.String, primary_key=True, server_default=FetchedValue())
    form_spec_json = db.Column(db.String, nullable=False)
    source_model_name = db.Column(db.String, nullable=False)
    template_file_path = db.Column(db.String, nullable=False)
    active_ind = db.Column(db.String, nullable=False, server_default=FetchedValue())

    context_primary_key = None

    def __repr__(self):
        return '<DocumentTemplate %r>' % self.document_template_code

    @hybrid_property
    def form_spec(self):
        if self.context_primary_key:
            return self._form_spec_with_context(self.context_primary_key)

        return json.loads(self.form_spec_json)

    def _form_spec_with_context(self, primary_key):
        spec = json.loads(self.form_spec_json)
        current_app.logger.debug(f'getting model for string -> {self.source_model_name}')

        source_model = get_model_by_model_name(self.source_model_name)
        current_app.logger.debug(f'source_obj_model -> {source_model}')

        source_obj_instance = source_model.query.get(primary_key)
        current_app.logger.debug(f'source_obj_instance -> {source_obj_instance}')
        if not source_obj_instance:
            raise Exception('Context Object not found')

        for item in spec:

            # Handle "relative data path"
            relative_data_path = item.get('relative-data-path')
            if relative_data_path:
                current_object = source_obj_instance
                for x in relative_data_path.split('.'):
                    current_app.logger.debug(f'getting {current_object}.{x}')
                    current_object = getattr(current_object, x)
                    if current_object is None:
                        break

                current_app.logger.debug(
                    f'Found data for form."{item["id"]}" at "{item["relative-data-path"]}" with -> "{current_object}"'
                )

                del item['relative-data-path']
                context_value = str(current_object) if current_object else None
                item['context-value'] = context_value
                continue

            # Handle "context-value"
            context_value = item.get('context-value')
            if context_value:
                if context_value == '{DATETIME.UTCNOW}':
                    context_value = str(datetime.utcnow())

                item['context-value'] = context_value
                continue

        return spec

    @hybrid_property
    def os_template_file_path(self):
        return os.path.join(current_app.root_path, self.template_file_path)

    @hybrid_property
    def template_name(self):
        return self.template_file_path.split('/')[-1]

    @hybrid_property
    def template_name_no_extension(self):
        return '.'.join(self.template_name.split('.')[:-1])

    def get_dynamic_template(self, template_data):
        def insert_images(doc, template_data):
            images = template_data.get('images', {})
            for key in images:
                image = images[key]
                image_base64 = image['source']
                if not image_base64:
                    continue

                for paragraph in doc.paragraphs:
                    if key in paragraph.text:
                        image_data = None
                        try:
                            image_data = base64.b64decode(image_base64.split(',')[1])
                        except:
                            raise Exception('Image data is not a valid Base64 string')

                        image_bytes = io.BytesIO(image_data)
                        width = Inches(image['width']) if image['width'] else None
                        height = Inches(image['height']) if image['height'] else None
                        paragraph.clear()
                        run = paragraph.add_run()
                        run.add_picture(image_bytes, width=width, height=height)

            if template_data.get('images'):
                del template_data['images']

        def replace_html(self, template_data):
            if template_data['letter_body']:
                for paragraph in doc.paragraphs:
                    if 'letter_body' in paragraph.text:
                        paragraph.clear()
                        #soup = BeautifulSoup(template_data['letter_body'], 'html.parser')
                        #for paragraph in soup.find_all('p'):
                        #    run = paragraph.add_run('-------------------TEST---------------')
                        html = template_data['letter_body'].split('<')
                        html = [html[0]] + ['<' + l for l in html[1:]]
                        for run in html:
                            if run.startswith('<p>'):
                                run = run.lstrip('<p>')
                            if run.startswith('</p>'):
                                run = run.lstrip('</p>')
                            if run.startswith('<strong>'):
                                run = run.lstrip('<strong>')
                                runner = paragraph.add_run(run)
                                runner.bold = True
                            elif run.startswith('</strong>'):
                                run = run.lstrip('</strong>')
                                runner = paragraph.add_run(run)
                            else:
                                paragraph.add_run(run)

                del template_data['letter_body']

        doc = None
        if self.document_template_code in ('PMT', 'PMA', 'NPE', 'NCL', 'NWL', 'NRL'):
            doc = docx.Document(self.os_template_file_path)
            insert_images(doc, template_data)
            replace_html(doc, template_data)

        if doc:
            fileobj = io.BytesIO()
            doc.save(fileobj)
            fileobj.seek(0)
            return fileobj

        return None
